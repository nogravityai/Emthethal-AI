"""
core/extractor.py — Emthethal AI
The File Router: Routes files to the correct parser based on type.

RULES:
- DOCX / XLSX: Parse using native libraries (python-docx, openpyxl).
  Preserve native tables and cells directly. NO OCR or geometric clustering.
- PDF (Text-based & Scanned): Route to PaddleOCR (Arabic support) followed
  by the Geometry Engine for layout reconstruction.
- Images: Route to PaddleOCR directly.
"""

import io
import logging
from pathlib import Path
from typing import List, Optional, Tuple

import fitz  # PyMuPDF
import numpy as np

from ..ingestion_models.schemas import (
    ExtractedCell,
    TableRow,
    StructureBlock,
    PageOutput,
    DocumentOutput,
)
from .geometry_engine import GeometryEngine, OCRWord, geometry_engine

logger = logging.getLogger(__name__)


# ─── Lazy-Loaded PaddleOCR ────────────────────────────────────────────────────

_PADDLE_OCR = None


def _get_paddle_ocr():
    """
    Lazy-load PaddleOCR with Arabic + English support.

    CPU-safe initialization: sets CUDA_VISIBLE_DEVICES="" before importing
    paddle to prevent a Segmentation Fault that occurs when paddlepaddle-gpu
    is installed but CUDA runtime libraries are absent in the container.
    GPU mode can be re-enabled once CUDA libs are confirmed in the image.
    """
    global _PADDLE_OCR
    if _PADDLE_OCR is None:
        try:
            import os
            if os.environ.get("DISABLE_PADDLE"):
                _PADDLE_OCR = "unavailable"
                return _PADDLE_OCR
            os.environ["CUDA_VISIBLE_DEVICES"] = "-1"
            from paddleocr import PaddleOCR
            logger.info("🚀 Initializing PaddleOCR (lang=ar)...")
            _PADDLE_OCR = PaddleOCR(lang="ar")
            logger.info("✅ PaddleOCR initialized successfully")
        except ImportError:
            logger.warning(
                "PaddleOCR not available. Falling back to PyMuPDF text extraction. "
                "Install with: pip install paddlepaddle paddleocr"
            )
            _PADDLE_OCR = "unavailable"
        except Exception as e:
            logger.error(f"PaddleOCR initialization failed: {e}")
            _PADDLE_OCR = "unavailable"
    return _PADDLE_OCR


# ─── File Router ──────────────────────────────────────────────────────────────

class FileExtractor:
    """
    Routes files to the correct extraction strategy.
    Returns a DocumentOutput validated through Pydantic V2.
    """

    def __init__(self, geometry: Optional[GeometryEngine] = None):
        self.geometry = geometry or geometry_engine

    async def extract(
        self, file_path: Path, file_bytes: Optional[bytes] = None
    ) -> DocumentOutput:
        """
        Main entry point. Routes to the correct parser based on file extension.
        """
        ext = file_path.suffix.lower()

        if file_bytes is None:
            file_bytes = file_path.read_bytes()

        if ext == ".pdf":
            return self._extract_pdf(file_path, file_bytes)
        elif ext == ".docx":
            return self._extract_docx(file_path, file_bytes)
        elif ext in (".xlsx", ".xls"):
            return self._extract_xlsx(file_path, file_bytes)
        elif ext == ".doc":
            return self._extract_doc(file_path, file_bytes)
        elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
            return self._extract_image(file_path, file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    # ─── PDF Extraction (PaddleOCR + Geometry Engine) ─────────────────────────

    def _extract_pdf(self, file_path: Path, file_bytes: bytes) -> DocumentOutput:
        """
        PDF extraction pipeline:
        1. Try native text extraction first (for text-based PDFs)
        2. For each page, run PaddleOCR to get bboxes + confidence
        3. Route OCR output through Geometry Engine for layout reconstruction
        4. Fall back to raw text extraction if OCR unavailable
        """
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages: List[PageOutput] = []

        paddle = _get_paddle_ocr()
        use_paddle = paddle != "unavailable" and paddle is not None

        for page_idx in range(len(doc)):
            page = doc[page_idx]
            page_number = page_idx + 1

            if use_paddle:
                # Render page to image for OCR
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")
                img_array = np.frombuffer(img_bytes, dtype=np.uint8)

                # Use PIL to convert PNG bytes to numpy array
                from PIL import Image
                img = Image.open(io.BytesIO(img_bytes))
                img_np = np.array(img)

                # Run PaddleOCR (cls omitted — v3.x compatibility)
                try:
                    results = paddle.ocr(img_np)
                    ocr_words = self._paddle_results_to_words(results)

                    if ocr_words:
                        # Route through Geometry Engine
                        tables = self.geometry.reconstruct_layout(ocr_words)
                        block_dicts = self.geometry.tables_to_structure_blocks(tables)

                        blocks = []
                        for bd in block_dicts:
                            try:
                                blocks.append(StructureBlock(**bd))
                            except Exception as e:
                                logger.warning(f"Block validation failed on page {page_number}: {e}")

                        if blocks:
                            pages.append(PageOutput(
                                page_number=page_number,
                                blocks=blocks,
                            ))
                            continue
                except Exception as e:
                    logger.warning(f"PaddleOCR failed on page {page_number}: {e}")

            # Fallback: native text extraction
            raw_text = page.get_text().strip()
            if raw_text:
                # Create a single "mixed" block from raw text
                lines = [line.strip() for line in raw_text.split("\n") if line.strip()]
                rows = []
                for line in lines:
                    rows.append(TableRow(cells=[
                        ExtractedCell(text=line, confidence=0.85)
                    ]))

                if rows:
                    block = StructureBlock(
                        type="mixed",
                        confidence=0.85,
                        rows=rows,
                    )
                    pages.append(PageOutput(
                        page_number=page_number,
                        blocks=[block],
                        raw_text=raw_text,
                    ))

        doc.close()

        # If no pages extracted, create empty document
        if not pages:
            pages.append(PageOutput(page_number=1, blocks=[]))

        return DocumentOutput(
            document_type="medical_form",
            source_filename=file_path.name,
            file_type="pdf",
            pages=pages,
            processing_metadata={
                "ocr_engine": "paddleocr" if use_paddle else "pymupdf_text",
                "total_pages": len(pages),
            },
        )

    def _paddle_results_to_words(self, results) -> List[OCRWord]:
        """
        Convert PaddleOCR output to OCRWord list.

        Handles both PaddleOCR v2.x and v3.x output formats:
          v2.x: [ [ [bbox_points, (text, conf)], ... ] ]  — list of pages
          v3.x: same structure but None entries on blank pages
        """
        words = []
        if not results:
            return words

        for result_page in results:
            # v3.x returns None for blank/empty pages
            if result_page is None:
                continue
            if not result_page:
                continue

            for item in result_page:
                # Skip None or malformed items
                if item is None:
                    continue
                if not isinstance(item, (list, tuple)) or len(item) < 2:
                    continue

                bbox_points = item[0]  # [[x1,y1], [x2,y1], [x2,y2], [x1,y2]]
                text_conf = item[1]    # (text, confidence)

                if not bbox_points or not text_conf:
                    continue

                # text_conf can be tuple or list: (text, score)
                try:
                    text = str(text_conf[0]).strip()
                    confidence = float(text_conf[1]) if len(text_conf) > 1 else 0.5
                except (IndexError, TypeError, ValueError):
                    continue

                if not text:
                    continue

                # Convert 4-point polygon → axis-aligned [x1, y1, x2, y2]
                try:
                    xs = [float(p[0]) for p in bbox_points]
                    ys = [float(p[1]) for p in bbox_points]
                    bbox = [min(xs), min(ys), max(xs), max(ys)]
                except (IndexError, TypeError, ValueError):
                    continue

                words.append(OCRWord(
                    text=text,
                    bbox=bbox,
                    confidence=confidence,
                ))

        return words

    # ─── DOCX Extraction (Native Parser) ──────────────────────────────────────

    def _extract_docx(self, file_path: Path, file_bytes: bytes) -> DocumentOutput:
        """
        DOCX extraction using python-docx.
        Preserves native tables with cell-level accuracy.
        NO OCR, NO geometric clustering.
        """
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        pages: List[PageOutput] = []
        blocks: List[StructureBlock] = []

        # Extract tables (primary structured data)
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    cells.append(ExtractedCell(
                        text=cell_text,
                        confidence=1.0,  # Native parsing = perfect confidence
                        bbox=None,  # No bbox for DOCX
                    ))
                if any(c.text for c in cells):
                    rows.append(TableRow(cells=cells))

            if rows:
                # Detect headers from first row
                headers = [c.text for c in rows[0].cells] if rows else None

                blocks.append(StructureBlock(
                    type="table",
                    confidence=1.0,
                    rows=rows,
                    headers=headers,
                    title=f"Table {table_idx + 1}",
                ))

        # Extract paragraphs as form blocks
        paragraph_rows = []
        current_section = None
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style = para.style.name
            if style.startswith("Heading"):
                # Flush accumulated paragraphs
                if paragraph_rows:
                    blocks.append(StructureBlock(
                        type="form",
                        confidence=1.0,
                        rows=paragraph_rows,
                        title=current_section,
                    ))
                    paragraph_rows = []
                current_section = text
            else:
                paragraph_rows.append(TableRow(cells=[
                    ExtractedCell(text=text, confidence=1.0)
                ]))

        # Flush remaining paragraphs
        if paragraph_rows:
            blocks.append(StructureBlock(
                type="form",
                confidence=1.0,
                rows=paragraph_rows,
                title=current_section,
            ))

        if blocks:
            pages.append(PageOutput(page_number=1, blocks=blocks))
        else:
            pages.append(PageOutput(page_number=1, blocks=[]))

        return DocumentOutput(
            document_type="medical_form",
            source_filename=file_path.name,
            file_type="docx",
            pages=pages,
            processing_metadata={
                "parser": "python-docx",
                "tables_found": sum(1 for b in blocks if b.type == "table"),
                "form_sections_found": sum(1 for b in blocks if b.type == "form"),
            },
        )

    # ─── XLSX Extraction (Native Parser) ──────────────────────────────────────

    def _extract_xlsx(self, file_path: Path, file_bytes: bytes) -> DocumentOutput:
        """
        XLSX extraction using openpyxl.
        Each worksheet becomes a page, each contiguous data region a block.
        NO OCR, NO geometric clustering.
        """
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
        pages: List[PageOutput] = []

        for sheet_idx, sheet_name in enumerate(wb.sheetnames):
            ws = wb[sheet_name]
            blocks: List[StructureBlock] = []

            # Read all rows from the worksheet
            rows_data = []
            for row in ws.iter_rows(values_only=False):
                cells = []
                has_data = False
                for cell in row:
                    val = cell.value
                    text = str(val).strip() if val is not None else ""
                    if text:
                        has_data = True
                    cells.append(ExtractedCell(
                        text=text,
                        value=text if text else None,
                        confidence=1.0,  # Native parsing = perfect confidence
                        bbox=None,
                    ))

                if has_data:
                    rows_data.append(TableRow(cells=cells))

            if rows_data:
                # Detect headers from first row
                headers = [c.text for c in rows_data[0].cells] if rows_data else None

                blocks.append(StructureBlock(
                    type="table",
                    confidence=1.0,
                    rows=rows_data,
                    headers=headers,
                    title=sheet_name,
                ))

            pages.append(PageOutput(
                page_number=sheet_idx + 1,
                blocks=blocks,
            ))

        wb.close()

        return DocumentOutput(
            document_type="medical_form",
            source_filename=file_path.name,
            file_type="xlsx",
            pages=pages,
            processing_metadata={
                "parser": "openpyxl",
                "sheets": len(pages),
            },
        )

    # ─── DOC Extraction (Legacy) ──────────────────────────────────────────────

    def _extract_doc(self, file_path: Path, file_bytes: bytes) -> DocumentOutput:
        """Legacy .doc extraction via docx2txt."""
        import docx2txt
        import tempfile
        import os

        # docx2txt needs a file path, write temp file
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            text = docx2txt.process(tmp_path) or ""
        finally:
            os.unlink(tmp_path)

        rows = []
        for line in text.split("\n"):
            line = line.strip()
            if line:
                rows.append(TableRow(cells=[
                    ExtractedCell(text=line, confidence=0.7)
                ]))

        blocks = []
        if rows:
            blocks.append(StructureBlock(
                type="mixed",
                confidence=0.7,
                rows=rows,
            ))

        return DocumentOutput(
            document_type="medical_form",
            source_filename=file_path.name,
            file_type="doc",
            pages=[PageOutput(page_number=1, blocks=blocks)],
            processing_metadata={"parser": "docx2txt"},
        )

    # ─── Image Extraction (PaddleOCR) ─────────────────────────────────────────

    def _extract_image(self, file_path: Path, file_bytes: bytes) -> DocumentOutput:
        """Direct image OCR via PaddleOCR + Geometry Engine."""
        from PIL import Image

        img = Image.open(io.BytesIO(file_bytes))
        img_np = np.array(img)

        paddle = _get_paddle_ocr()
        blocks: List[StructureBlock] = []

        if paddle != "unavailable" and paddle is not None:
            try:
                results = paddle.ocr(img_np)
                ocr_words = self._paddle_results_to_words(results)

                if ocr_words:
                    tables = self.geometry.reconstruct_layout(ocr_words)
                    block_dicts = self.geometry.tables_to_structure_blocks(tables)

                    for bd in block_dicts:
                        try:
                            blocks.append(StructureBlock(**bd))
                        except Exception as e:
                            logger.warning(f"Block validation failed for image: {e}")
            except Exception as e:
                logger.error(f"Image OCR failed: {e}")

        return DocumentOutput(
            document_type="medical_form",
            source_filename=file_path.name,
            file_type="image",
            pages=[PageOutput(page_number=1, blocks=blocks)],
            processing_metadata={"parser": "paddleocr"},
        )


# ─── Module-Level Instance ────────────────────────────────────────────────────

file_extractor = FileExtractor()
