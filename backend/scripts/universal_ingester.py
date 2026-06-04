import os
import sys
import logging
import asyncio
import json
import re
import io
from pathlib import Path
from datetime import datetime

# OCR and Image Processing
import easyocr
import numpy as np
from PIL import Image
import fitz  # PyMuPDF

sys.path.append(os.path.join(os.path.dirname(__file__), '..'))

import httpx
from app.database import async_session
from app.models import DocumentChunk

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

OLLAMA_BASE_URL = os.getenv("OLLAMA_URL", "http://ollama:11434")
EMBED_MODEL = "nomic-embed-text"

QUEUE_DIR = Path(__file__).parent.parent / "ingestion_queue"
PROCESSED_DIR = QUEUE_DIR / "processed"
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".png", ".jpg", ".jpeg"}
MAX_CHARS = 2000

# Global OCR Reader (Lazy Loaded)
_READER = None

def get_reader():
    global _READER
    if _READER is None:
        logger.info("🚀 Initializing EasyOCR with Arabic + English support (GPU enabled)...")
        _READER = easyocr.Reader(['ar', 'en'], gpu=True)
    return _READER

# ─── Embedding ────────────────────────────────────────────────────────────────

async def get_embedding(text: str) -> list[float]:
    """Get vector embedding from Ollama/LM Studio."""
    base_url = OLLAMA_BASE_URL.rstrip('/')
    is_openai = "/v1" in base_url or "1234" in base_url
    
    if is_openai:
        url = f"{base_url}/embeddings" if "/v1" in base_url else f"{base_url}/v1/embeddings"
        payload = {"model": EMBED_MODEL, "input": text[:MAX_CHARS]}
    else:
        url = f"{base_url}/api/embeddings"
        payload = {"model": EMBED_MODEL, "prompt": text[:MAX_CHARS]}
        
    async with httpx.AsyncClient(timeout=60.0) as client:
        for attempt in range(3):
            try:
                response = await client.post(url, json=payload)
                response.raise_for_status()
                res_data = response.json()
                if is_openai:
                    return res_data["data"][0]["embedding"]
                else:
                    return res_data["embedding"]
            except Exception as e:
                if attempt == 2: raise e
                await asyncio.sleep(1)

# ─── OCR Helper ───────────────────────────────────────────────────────────────

def ocr_image(image_bytes: bytes) -> str:
    """Perform GPU-accelerated OCR on image bytes."""
    try:
        reader = get_reader()
        # EasyOCR can take bytes, numpy array, or path
        result = reader.readtext(image_bytes, detail=0)
        return " ".join(result)
    except Exception as e:
        logger.error(f"  ❌ OCR Error: {e}")
        return ""

# ─── Parsers ──────────────────────────────────────────────────────────────────

def parse_pdf(file_path: Path) -> list[dict]:
    """Extract text and images from PDF. Each image becomes its own chunk."""
    doc = fitz.open(str(file_path))
    chunks = []
    
    # 1. Text Extraction
    full_text = ""
    for page in doc:
        full_text += page.get_text()
    
    if full_text.strip():
        chunks.append({
            "content": full_text,
            "department": "General",
            "device_name": "Digital Text Content"
        })

    # 2. Image Extraction & OCR
    for page_index in range(len(doc)):
        page = doc[page_index]
        image_list = page.get_images(full=True)
        
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            
            logger.info(f"    - OCR-ing Image {img_index+1} on Page {page_index+1}...")
            text = ocr_image(image_bytes)
            
            if text.strip():
                chunks.append({
                    "content": text,
                    "department": "General",
                    "device_name": f"Image Content (Page {page_index+1}, Img {img_index+1})"
                })
    
    doc.close()
    return chunks

def parse_docx(file_path: Path) -> list[dict]:
    """Extract text, tables, and images from DOCX."""
    from docx import Document
    doc = Document(str(file_path))
    chunks, current_dept, current_device, buf = [], "General", "General", []

    # 1. Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        style = para.style.name
        if style.startswith("Heading 1"):
            current_dept = text
        elif style.startswith("Heading 2"):
            current_device = text
        else:
            buf.append(text)
    
    if buf:
        chunks.append({"content": "\n".join(buf), "department": current_dept, "device_name": current_device})

    # 2. Tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            table_data.append(" | ".join(filter(None, cells)))
        if table_data:
            chunks.append({
                "content": "\n".join(table_data),
                "department": current_dept,
                "device_name": "Table Content"
            })

    # 3. Images
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_bytes = rel.target_part.blob
                logger.info("    - OCR-ing embedded image in DOCX...")
                text = ocr_image(img_bytes)
                if text.strip():
                    chunks.append({
                        "content": text,
                        "department": current_dept,
                        "device_name": "Embedded Image Content"
                    })
    except Exception as e:
        logger.warning(f"    ⚠️ DOCX Image OCR failed: {e}")

    return chunks

def parse_xlsx(file_path: Path) -> list[dict]:
    """Simple XLSX parser."""
    import pandas as pd
    df = pd.read_excel(file_path)
    content = df.to_string()
    return [{"content": content, "department": "General", "device_name": "Excel Content"}]

def parse_doc(file_path: Path) -> list[dict]:
    """Legacy .doc support via docx2txt."""
    import docx2txt
    text = docx2txt.process(str(file_path))
    return [{"content": text, "department": "General", "device_name": "Legacy Document"}]

# ─── File processor ───────────────────────────────────────────────────────────

async def process_file(file_path: Path, relative_folder: str):
    ext = file_path.suffix.lower()
    form_type = relative_folder.strip("/") or "General"
    logger.info(f"\n📄 {file_path.name}  |  form_type={form_type}  ({ext.upper()})")

    try:
        if ext == ".pdf":
            chunks = parse_pdf(file_path)
        elif ext == ".docx":
            chunks = parse_docx(file_path)
        elif ext == ".doc":
            chunks = parse_doc(file_path)
        elif ext in (".xlsx", ".xls"):
            chunks = parse_xlsx(file_path)
        elif ext in (".png", ".jpg", ".jpeg"):
            with open(file_path, "rb") as f:
                text = ocr_image(f.read())
            chunks = [{"content": text, "department": "General", "device_name": "Stand-alone Image"}] if text.strip() else []
        else:
            logger.warning(f"  ⚠️ Unsupported: {ext}")
            return

        if not chunks:
            logger.warning(f"  ⚠️ No content extracted from {file_path.name}")
            return

        # Commit chunks to DB
        # Commit chunks to DB
        async with async_session() as session:
            count = 0
            
            # Split form_type (which is the relative path) into levels
            path_parts = relative_folder.strip("/").split("/")
            
            # Level 1: Department
            dept = path_parts[0] if len(path_parts) > 0 else "General"
            
            # Level 2: Device / Asset (Subfolder)
            device_attr = path_parts[1] if len(path_parts) > 1 else "General"
            
            # Level 3: Specific Form (Filename or sub-subfolder)
            # If the chunk has a specific title from parsing (like a table title), use it.
            # Otherwise, use the filename.
            
            for chunk in chunks:
                final_device = chunk.get("device_name", "General")
                if final_device == "General":
                    final_device = device_attr
                
                metadata = {
                    "department": dept,
                    "device_name": final_device,
                    "form_type": file_path.stem, # Use filename without extension as form type
                    "source": file_path.name,
                    "path": str(relative_folder),
                    "ingested_at": datetime.now().isoformat()
                }

                emb = await get_embedding(chunk["content"])
                db_chunk = DocumentChunk(
                    content=chunk["content"],
                    embedding=emb,
                    metadata_payload=metadata
                )
                session.add(db_chunk)
                count += 1
            
            await session.commit()
            logger.info(f"  💾 {count}/{len(chunks)} chunks committed.")

        # Move to processed
        dest = PROCESSED_DIR / relative_folder / file_path.name
        dest.parent.mkdir(parents=True, exist_ok=True)
        file_path.rename(dest)
        logger.info(f"  📦 → {dest.relative_to(QUEUE_DIR)}")

    except Exception as e:
        logger.error(f"  ❌ Error processing {file_path.name}: {e}")

# ─── Main loop ────────────────────────────────────────────────────────────────

async def main():
    if not QUEUE_DIR.exists():
        QUEUE_DIR.mkdir(parents=True)
    
    logger.info("🚀 Starting Universal Ingester with EasyOCR (GPU)...")
    
    # Process files recursively
    files_to_process = []
    for ext in SUPPORTED_EXTENSIONS:
        files_to_process.extend(list(QUEUE_DIR.rglob(f"*{ext}")))
    
    # Filter out files already in processed/
    files_to_process = [f for f in files_to_process if "processed" not in str(f)]

    if not files_to_process:
        logger.info("✅ No new files to process.")
        return

    logger.info(f"🔍 Found {len(files_to_process)} files to process.")

    for file_path in files_to_process:
        rel_folder = str(file_path.parent.relative_to(QUEUE_DIR))
        await process_file(file_path, rel_folder)

if __name__ == "__main__":
    asyncio.run(main())
